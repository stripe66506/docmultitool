# DocSpeak program- convert text files into audio

#reqiured libraries
import pyttsx3
import os
import docx2txt
from pyth.plugins.rtf15.reader import Rtf15Reader
from pdfminer.high_level import extract_text
import PyPDF2
import re
import pytesseract
from pdf2image import convert_from_path
from rake_nltk import Rake
from docx import Document

def main_menu():
    print("\n--- Welcome to DocMultiTool! ---\n")
    print("Please select an option:")
    print("1. Convert PDF to TXT")
    print("2. Convert text file to Audio")
    print("3. Keyword Extractor")
    print("4. Feedback Collector")
    print("5. Exit")

    choice = input("\nEnter your choice: ")

    if choice == '1':
        print("You've selected 'Convert PDF to TXT' option.")
        pdf2doc()

    elif choice == '2':
        print("You've selected 'Convert text file to Audio' option.")
        run_docspeak()

    elif choice == '3':
        print("You've selected the 'Keyword Extractor' option.")
        extract_keywords()

    elif choice == '4':
        print("You've selected 'Feedback Selector")
        feedback_collection()

    elif choice == '5':
        print("Exiting program. Goodbye!")
        quit()

    else:
        print("Invalid option. Please try again.")
        main_menu()  # Prompt the user again if input was invalid

def feedback_collection():
    doc = Document()
    feedbacks = []
    while True:
        feedback = input("Enter your feedback (type 'quit' to stop): ")
        if feedback == 'quit':
            break
        feedbacks.append(feedback)
        doc.add_paragraph(feedback)

    # Create the subfolder if it doesn't exist
    folder_path = 'Feedbacks'  # Replace with the desired folder name
    if not os.path.exists(folder_path):
        os.makedirs(folder_path)

    # Save feedbacks to a text file within the subfolder
    file_path = os.path.join(folder_path, 'feedbacks.txt')
    with open(file_path, 'w') as file:
        file.write('\n'.join(feedbacks))

    # Save the document file within the subfolder
    doc_file_path = os.path.join(folder_path, 'feedbacks.docx')
    doc.save(doc_file_path)

    while True:
        d = input('Do you want to view feedbacks or rerun function (y or n): ')
        if d in ['y', 'n']:
            break
        else:
            print('Invalid input. Please type \'y\' or \'n\'. ')
    if d == 'y':
        feedback_collection()
    else:
        main_menu()

def extract_keywords():
    file_path = input('Please enter the file path: ')

    # Check file type by extension
    if file_path.endswith('.txt'):
        with open(file_path, 'r', encoding='utf8') as file:
            text = file.read()

    elif file_path.endswith('.docx'):
        doc = Document(file_path)
        text = ' '.join([para.text for para in doc.paragraphs])

    elif file_path.endswith('.pdf'):
        pdf_file_obj = open(file_path, 'rb', encoding='utf8')
        pdf_reader = PyPDF2.PdfFileReader(pdf_file_obj)
        text = ''
        for page_num in range(pdf_reader.numPages):
            page_obj = pdf_reader.getPage(page_num)
            text += page_obj.extractText()
        pdf_file_obj.close()

    else:
        print('File type not supported.')
        return

    # extract keywords with Rake
    r = Rake()
    r.extract_keywords_from_text(text)
    keywords = r.get_ranked_phrases()  # To get keyword phrases ranked highest to lowest.
    keywords_str =  '\n'.join(keywords)

    txt_path = input('What would you like to name the export file?: ') + '.txt'
    with open(txt_path, 'w', encoding='utf-8') as txt_file:  # specify 'utf-8' encoding
        txt_file.write(keywords_str)

    while True:
        d = input('Do you want to find keywords of another document (y or n): ')

        if d in ['y', 'n']:
            break

        else:
            print('Invalid input. Please type \'y\' or \'n\'. ')

    if d in ['y']:
        extract_keywords()

    else:
        main_menu()

def pdf2doc():
    c = input('Type the file name of your pdf: ')
    pdf_path = c + '.pdf'

    with open(pdf_path, 'rb') as pdf_file:
        pdf_reader = PyPDF2.PdfReader(pdf_file)

        # Initialize an empty string to hold the PDF text
        pdf_text = ''

        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            pdf_text += page.extract_text()

        pdf_text = re.sub(r'\n[0-9]+\.\s', '\n', pdf_text)

    txt_path = input('Type the name of your text file: ') + '.txt'
    with open(txt_path, 'w', encoding='utf-8') as txt_file:  # specify 'utf-8' encoding
        txt_file.write(pdf_text)

    while True:
        d = input('Do you want to convert another PDF (y or n): ')

        if d in ['y', 'n']:
            break

        else:
            print('Invalid input. Please type \'y\' or \'n\'. ')

    if d in ['y']:
        pdf2doc()

    else:
        main_menu()

def run_docspeak():

    input_file = input('Enter name of your text file: ')
    a = input('Enter name of the desired mp3: ')
    output_file = a + '.mp3'

    # Text from the input file
    if input_file.endswith('.docx'):
        text = docx2txt.process(input_file)
    elif input_file.endswith('.txt'):
        with open(input_file,'r', encoding='utf-8') as file:
            text = file.read()
    elif input_file.endswith('.rtf'):
        with open(input_file,'rb', encoding='utf-8') as file:
            doc = Rtf15Reader.read(file)
            text = ''.join([paragraph.text for paragraph in doc.content if hasattr(paragraph, 'text')])
    elif input_file.endswith('.pdf'):
        text = extract_text(input_file)
    else:
        print('Unsupported file type. Please use .txt, .docx, .rtf or .pdf file.')
        return  # Return to avoid running the rest of the function

    # Initialize Text To Speech engine
    engine = pyttsx3.init()

    # Save the audio as mp3
    engine.save_to_file(text, output_file)
    engine.runAndWait()

    while True:
        d = input('Do you want to convert another text file? (y or n): ')

        if d in ['y', 'n']:
            break
        else:
            print('Invalid input. Please type \'y\' or \'n\'. ')

    if d == 'y':
        run_docspeak()
    else:
        main_menu()

def pdf_to_word():
    # Ask for input and output file names
    a = input('Enter name of the pdf file')
    pdf_file = a + '.pdf'
    word_file = input('Enter the name of word file')

    # convert pdf into image
    pages = convert_from_path(pdf_file)

    # Create a new word document
    doc = Document()

    # Perform OCR on each page and add the text to the Word document
    for page in pages:
        text = pytesseract.image_to_string(page)
        doc.add_paragraph(text)

    # Save the document

    doc.save(word_file)

# Run the main menu function
main_menu()






